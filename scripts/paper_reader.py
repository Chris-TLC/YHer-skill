#!/usr/bin/env python3
"""
Exam paper document reader: handles docx / pdf / .doc uniformly and outputs plain text.

- docx: python-docx, keeps paragraphs + tables
- pdf: PyMuPDF (fitz)
- .doc (legacy format): convert with macOS textutil; if conversion fails (scanned doc), return empty and flag it
"""

from __future__ import annotations

import subprocess
from pathlib import Path
from typing import Tuple


def read_paper(path: Path) -> Tuple[str, str]:
    """
    Read an exam paper, return (text, status).
    status: ok / empty / unreadable / unsupported
    """
    path = Path(path)
    ext = path.suffix.lower()
    try:
        if ext == ".docx":
            return _read_docx(path)
        if ext == ".pdf":
            return _read_pdf(path)
        if ext == ".doc":
            return _read_doc(path)
        return "", "unsupported"
    except Exception as e:
        return f"[read failed: {e}]", "unreadable"


def _read_docx(path: Path) -> Tuple[str, str]:
    import docx
    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))  # separate table cells with | to preserve structure
    text = "\n".join(x for x in parts if x.strip())
    return text, ("ok" if len(text) > 200 else "empty")


def _read_pdf(path: Path) -> Tuple[str, str]:
    import fitz
    doc = fitz.open(str(path))
    text = "\n".join(page.get_text() for page in doc)
    return text, ("ok" if len(text) > 200 else "empty")


def _read_doc(path: Path) -> Tuple[str, str]:
    """Convert legacy .doc with macOS textutil. Scanned docs that fail conversion return empty."""
    try:
        out = subprocess.run(
            ["textutil", "-convert", "txt", "-stdout", str(path)],
            capture_output=True, timeout=30,
        )
        text = out.stdout.decode("utf-8", errors="ignore")
        return text, ("ok" if len(text) > 200 else "empty")
    except FileNotFoundError:
        return "", "unsupported"  # non-macOS environment
    except Exception as e:
        return f"[doc conversion failed: {e}]", "unreadable"


if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        text, status = read_paper(Path(sys.argv[1]))
        print(f"status={status}, chars={len(text)}")
        print(text[:800])
